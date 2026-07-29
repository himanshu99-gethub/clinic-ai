"""
Document Loader — supports PDF, DOCX, TXT, CSV, XLSX
Returns a list of dicts: {filename, content, metadata}
"""

import io
import os
from typing import List, Dict, Any

import pandas as pd


def load_document(filename: str, file_bytes: bytes) -> List[Dict[str, Any]]:
    """
    Load a document from raw bytes and extract text content.
    Returns a list of page/chunk dicts with 'content' and 'metadata'.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _load_pdf(filename, file_bytes)
    elif ext == ".docx":
        return _load_docx(filename, file_bytes)
    elif ext == ".txt":
        return _load_txt(filename, file_bytes)
    elif ext == ".csv":
        return _load_csv(filename, file_bytes)
    elif ext in (".xlsx", ".xls"):
        return _load_xlsx(filename, file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _load_pdf(filename: str, file_bytes: bytes) -> List[Dict[str, Any]]:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({
                    "content": text,
                    "metadata": {"source": filename, "page": i + 1}
                })
        return pages
    except Exception as e:
        return [{"content": "", "metadata": {"source": filename, "error": str(e)}}]


def _load_docx(filename: str, file_bytes: bytes) -> List[Dict[str, Any]]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text += "\n" + row_text
        return [{"content": full_text, "metadata": {"source": filename, "page": 1}}]
    except Exception as e:
        return [{"content": "", "metadata": {"source": filename, "error": str(e)}}]


def _load_txt(filename: str, file_bytes: bytes) -> List[Dict[str, Any]]:
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        return [{"content": text, "metadata": {"source": filename, "page": 1}}]
    except Exception as e:
        return [{"content": "", "metadata": {"source": filename, "error": str(e)}}]


def _load_csv(filename: str, file_bytes: bytes) -> List[Dict[str, Any]]:
    try:
        df = pd.read_csv(io.BytesIO(file_bytes))
        # Convert entire DataFrame to string for embedding
        text = df.to_string(index=False)
        # Also create comma-separated rows for email scanning
        raw_rows = "\n".join(
            ", ".join(str(v) for v in row.values) for _, row in df.iterrows()
        )
        return [{"content": text + "\n" + raw_rows, "metadata": {"source": filename, "rows": len(df)}}]
    except Exception as e:
        return [{"content": "", "metadata": {"source": filename, "error": str(e)}}]


def _load_xlsx(filename: str, file_bytes: bytes) -> List[Dict[str, Any]]:
    try:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        all_text = []
        for sheet_name, sheet_df in df.items():
            sheet_text = f"[Sheet: {sheet_name}]\n" + sheet_df.to_string(index=False)
            raw_rows = "\n".join(
                ", ".join(str(v) for v in row.values) for _, row in sheet_df.iterrows()
            )
            all_text.append(sheet_text + "\n" + raw_rows)
        return [{"content": "\n\n".join(all_text), "metadata": {"source": filename}}]
    except Exception as e:
        return [{"content": "", "metadata": {"source": filename, "error": str(e)}}]
